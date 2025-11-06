#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.enum.section import WD_SECTION

print("Создание документа требований по ГОСТу...")
sys.stdout.flush()

try:
    import docx
    print("Библиотека python-docx загружена успешно")
except ImportError:
    print("ОШИБКА: библиотека python-docx не установлена")
    sys.exit(1)

# Создаем новый документ
doc = Document()

# Настройка полей по ГОСТу (верхнее и нижнее - 2 см, левое - 3 см, правое - 1.5 см)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    
    # Добавляем нумерацию страниц в нижний колонтитул
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()  # Очищаем существующий текст
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем поле номера страницы (Times New Roman 14)
    page_field = parse_xml(
        '<w:fldSimple w:instr="PAGE" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman"/><w:sz w:val="28"/><w:color w:val="000000"/></w:rPr>'
        '<w:t>1</w:t></w:r></w:fldSimple>'
    )
    footer_para._element.append(page_field)

# Функция для настройки стиля по ГОСТу
def set_gost_style(paragraph, font_name='Times New Roman', font_size=14, bold=False, 
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5):
    """Настройка стиля параграфа по ГОСТу"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = line_spacing
    paragraph_format.first_line_indent = Cm(1.25)  # Отступ первой строки 1.25 см
    paragraph_format.alignment = alignment
    
    # Устанавливаем Times New Roman 14 и черный цвет для всех runs
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
        if bold:
            run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Функция для параграфа без отступа (для списков, заголовков)
def set_no_indent_style(paragraph, font_name='Times New Roman', font_size=14, bold=False):
    """Стиль без отступа первой строки"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Cm(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.5
    
    # Устанавливаем Times New Roman 14 и черный цвет для всех runs
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
        if bold:
            run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Функция для добавления закладки к параграфу
def add_bookmark(paragraph, bookmark_name):
    """Добавляет закладку к параграфу для создания гиперссылок"""
    # Получаем первый run или создаем новый
    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run('')
    
    # Создаем закладки
    bookmark_start = parse_xml(
        f'<w:bookmarkStart w:id="0" w:name="{bookmark_name}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    bookmark_end = parse_xml(
        f'<w:bookmarkEnd w:id="0" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    
    # Вставляем закладки в параграф
    paragraph._element.insert(0, bookmark_start)
    paragraph._element.append(bookmark_end)

# Функция для добавления гиперссылки
def add_hyperlink(paragraph, text, bookmark):
    """Добавляет гиперссылку на закладку"""
    part = paragraph.part
    r_id = part.relate_to(
        part.package,
        f'#{bookmark}',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=False
    )
    
    hyperlink = parse_xml(
        f'<w:hyperlink w:anchor="{bookmark}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0000FF"/><w:u w:val="single"/></w:rPr>'
        f'<w:t>{text}</w:t></w:r></w:hyperlink>'
    )
    
    paragraph._element.append(hyperlink)
    
    # Устанавливаем шрифт Times New Roman для гиперссылки
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# Функция для установки Times New Roman для всех runs в параграфе
def set_times_new_roman(paragraph):
    """Устанавливает Times New Roman 14 для всех runs в параграфе, черный цвет"""
    is_heading = paragraph.style.name.startswith('Heading')
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
        # Заголовки должны быть жирными по ГОСТу, остальной текст - нет
        if is_heading:
            run.bold = True
        else:
            run.bold = False
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# ========== ТИТУЛЬНЫЙ ЛИСТ ПО ГОСТУ ==========
# Пустые строки для центрирования
for _ in range(8):
    doc.add_paragraph()

# Название документа
title_para = doc.add_paragraph('ТРЕБОВАНИЯ К ПРОЕКТУ')
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
if title_para.runs:
    run = title_para.runs[0]
else:
    run = title_para.add_run()
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
title_para.paragraph_format.first_line_indent = Cm(0)
title_para.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# Название продукта
subtitle_para = doc.add_paragraph('Игра "Фэнтези Стратегия"')
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
if subtitle_para.runs:
    run = subtitle_para.runs[0]
else:
    run = subtitle_para.add_run()
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
subtitle_para.paragraph_format.first_line_indent = Cm(0)
subtitle_para.paragraph_format.line_spacing = 1.5

# Пустые строки перед подписью
for _ in range(10):
    doc.add_paragraph()

# Подпись (выравнивание по правому краю)
sign_para = doc.add_paragraph('2025')
sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_no_indent_style(sign_para, font_size=14)

# Разрыв страницы
doc.add_page_break()

# ========== СОДЕРЖАНИЕ ==========
content_title = doc.add_heading('СОДЕРЖАНИЕ', 1)
content_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_times_new_roman(content_title)

doc.add_paragraph()

# Список разделов для содержания с закладками
content_items = [
    ('1', 'Введение', 'bookmark_intro'),
    ('2', 'Требования пользователя', 'bookmark_user_req'),
    ('2.1', 'Программные интерфейсы', 'bookmark_interfaces'),
    ('2.2', 'Интерфейс пользователя', 'bookmark_ui'),
    ('2.3', 'Характеристики пользователей', 'bookmark_users'),
    ('2.4', 'Предположения и зависимости', 'bookmark_assumptions'),
    ('3', 'Системные требования', 'bookmark_system_req'),
    ('3.1', 'Функциональные требования', 'bookmark_func_req'),
    ('3.1.1', 'Описание системы заклинаний', 'bookmark_spells'),
    ('3.1.2', 'Описание системы юнитов', 'bookmark_units'),
    ('3.2', 'Нефункциональные требования', 'bookmark_nonfunc'),
    ('3.2.1', 'Атрибуты качества', 'bookmark_quality'),
    ('4', 'Заключение', 'bookmark_conclusion')
]

# Создаем содержание как список с гиперссылками и номерами страниц (черный текст, без подчеркивания, Times New Roman)
for num, name, bookmark in content_items:
    # Создаем параграф без маркеров
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    
    # Добавляем номер и название как гиперссылку (черный цвет, без подчеркивания, Times New Roman 14)
    run = p.add_run(f'{num} ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    run.bold = False  # Не жирный
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # Создаем гиперссылку на закладку (черный цвет, без подчеркивания, Times New Roman 14)
    hyperlink = parse_xml(
        f'<w:hyperlink w:anchor="{bookmark}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman"/><w:sz w:val="28"/><w:color w:val="000000"/><w:u w:val="none"/><w:b w:val="0"/></w:rPr>'
        f'<w:t>{name}</w:t></w:r></w:hyperlink>'
    )
    p._element.append(hyperlink)
    
    # Добавляем точки-заполнители и номер страницы с гиперссылкой
    # Используем поле PAGEREF для ссылки на страницу закладки
    dots_run = p.add_run(' ' + '.' * (50 - len(num) - len(name)) + ' ')
    dots_run.font.name = 'Times New Roman'
    dots_run.font.size = Pt(14)
    dots_run.font.color.rgb = RGBColor(0, 0, 0)
    dots_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # Добавляем поле PAGEREF для ссылки на страницу закладки (Times New Roman 14)
    # Это создаст гиперссылку на номер страницы раздела
    pageref_field = parse_xml(
        f'<w:fldSimple w:instr="PAGEREF {bookmark} \\h" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman"/><w:sz w:val="28"/><w:color w:val="000000"/></w:rPr>'
        '<w:t>1</w:t></w:r></w:fldSimple>'
    )
    p._element.append(pageref_field)
    
    # Устанавливаем Times New Roman для всех runs (не жирный)
    set_times_new_roman(p)

doc.add_page_break()

# ========== 1. ВВЕДЕНИЕ ==========
intro_heading = doc.add_heading('1 Введение', 1)
set_times_new_roman(intro_heading)
# Добавляем закладку для гиперссылки
add_bookmark(intro_heading, 'bookmark_intro')

intro_text = doc.add_paragraph()
intro_text.add_run('Название продукта: ')
intro_text.add_run('Фэнтези Стратегия')
set_gost_style(intro_text)

desc = doc.add_paragraph()
desc.add_run('Описание продукта: ')
desc.add_run('"Фэнтези Стратегия" — это пошаговая 2D стратегическая игра в фэнтези сеттинге, '
              'созданная на Python с использованием библиотеки Pygame. Игра представляет собой '
              'тактический боевой симулятор, в котором два игрока управляют героями и их армиями, '
              'сражаясь на поле боя с сеткой клеток. Каждый герой обладает уникальными способностями '
              'и заклинаниями, которые могут использоваться в бою.')
set_gost_style(desc)

boundaries = doc.add_paragraph()
boundaries.add_run('Границы проекта: ')
boundaries.add_run('Продукт представляет собой локальную игру для одного или двух игроков на одном устройстве. '
                   'Игра не включает сетевой мультиплеер, систему достижений, облачное сохранение прогресса, '
                   'или интеграцию с социальными сетями. Игра фокусируется на тактическом боевом геймплее '
                   'с пошаговым управлением юнитами и использованием магических заклинаний.')
set_gost_style(boundaries)

# ========== 2. ТРЕБОВАНИЯ ПОЛЬЗОВАТЕЛЯ ==========
user_req_heading = doc.add_heading('2 Требования пользователя', 1)
set_times_new_roman(user_req_heading)
add_bookmark(user_req_heading, 'bookmark_user_req')

# 2.1 Программные интерфейсы
interfaces_heading = doc.add_heading('2.1 Программные интерфейсы', 2)
set_times_new_roman(interfaces_heading)
add_bookmark(interfaces_heading, 'bookmark_interfaces')

interfaces = doc.add_paragraph('Продукт взаимодействует со следующими внешними системами и библиотеками:')
set_gost_style(interfaces)

interfaces_list = [
    'Pygame 2.5.2 — библиотека для создания игр и мультимедийных приложений, обеспечивает работу с графикой, звуком и вводом данных',
    'NumPy 1.24.3 — библиотека для численных вычислений, используется для математических операций в игровой логике',
    'Pillow >= 10.0.0 — библиотека для обработки изображений, используется для загрузки и обработки спрайтов и текстур',
    'Python 3.7+ — интерпретатор языка программирования, на котором реализована игра',
    'Операционная система Windows 10/11 — для нативной версии (EXE)',
    'Веб-браузер (Chrome, Firefox, Edge) — для веб-версии игры'
]

for item in interfaces_list:
    p = doc.add_paragraph(item, style='List Bullet')
    set_no_indent_style(p)

# 2.2 Интерфейс пользователя
ui_heading = doc.add_heading('2.2 Интерфейс пользователя', 2)
set_times_new_roman(ui_heading)
add_bookmark(ui_heading, 'bookmark_ui')

ui_desc = doc.add_paragraph('Интерфейс пользователя реализован в виде графического окна с разрешением 800x600 пикселей. '
                            'Игровое поле представлено сеткой клеток размером 40x40 пикселей каждая (20x15 клеток).')
set_gost_style(ui_desc)

doc.add_paragraph()

ui_table = doc.add_table(rows=1, cols=2)
ui_table.style = 'Light Grid Accent 1'
hdr_cells = ui_table.rows[0].cells
hdr_cells[0].text = 'Действие пользователя'
hdr_cells[1].text = 'Реакция системы'

# Стиль заголовков таблицы (не жирный, только Times New Roman)
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
            run.bold = False  # Не жирный
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)

actions = [
    ('Левая кнопка мыши на юните', 'Выбор юнита, подсветка доступных для перемещения клеток'),
    ('Левая кнопка мыши на пустой клетке', 'Перемещение выбранного юнита на указанную клетку (если доступно)'),
    ('Левая кнопка мыши на вражеском юните', 'Атака выбранным юнитом вражеского юнита (если в зоне атаки)'),
    ('Правая кнопка мыши на юните', 'Отображение тултипа с информацией о юните'),
    ('Двойной клик правой кнопкой на юните', 'Открытие окна с детальной информацией о юните'),
    ('Клавиши 1-3 (для героя)', 'Использование заклинаний героя (1 — первое заклинание, 2 — второе, 3 — третье)'),
    ('Клавиша ESC', 'Выход из текущего меню или режима'),
    ('Колесо мыши', 'Прокрутка интерфейса в режимах креатива и редактора книг'),
    ('F1-F6', 'Включение/выключение режимов отладки')
]

for action, reaction in actions:
    row_cells = ui_table.add_row().cells
    row_cells[0].text = action
    row_cells[1].text = reaction
    
    # Стиль для ячеек таблицы
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            paragraph.paragraph_format.first_line_indent = Cm(0)

# 2.3 Характеристики пользователей
users_heading = doc.add_heading('2.3 Характеристики пользователей', 2)
set_times_new_roman(users_heading)
add_bookmark(users_heading, 'bookmark_users')

users = doc.add_paragraph('Целевая аудитория игры:')
set_gost_style(users)

users_list = [
    'Игроки-любители стратегических игр — имеют базовый опыт в компьютерных играх, понимают основные механики пошаговых стратегий',
    'Казуальные игроки — могут иметь минимальный опыт в играх, но способны освоить простую систему управления',
    'Разработчики и модификаторы — имеют технические знания для работы с исходным кодом и модификации игры'
]

for item in users_list:
    p = doc.add_paragraph(item, style='List Bullet')
    set_no_indent_style(p)

tech_level = doc.add_paragraph()
tech_level.add_run('Техническая грамотность: ')
tech_level.add_run('Минимальная. Игра не требует специальных технических знаний для запуска и игры. '
                   'Для нативной версии (EXE) достаточно двойного клика по файлу. Для веб-версии требуется '
                   'умение открыть файл в браузере. Для работы с исходным кодом требуется знание Python.')
set_gost_style(tech_level)

# 2.4 Предположения и зависимости
assumptions_heading = doc.add_heading('2.4 Предположения и зависимости', 2)
set_times_new_roman(assumptions_heading)
add_bookmark(assumptions_heading, 'bookmark_assumptions')

assumptions = [
    'Игроки знакомы с базовыми концепциями пошаговых стратегий (ходы, атаки, перемещение)',
    'На устройстве установлены необходимые библиотеки Python (для версии из исходников) или отсутствуют (для EXE версии)',
    'Для веб-версии используется современный браузер с поддержкой WebAssembly',
    'Устройство имеет достаточную производительность для отображения 2D графики (минимальные требования: процессор 1 ГГц, 512 МБ RAM)',
    'Игроки понимают русский язык (интерфейс на русском языке)'
]

for item in assumptions:
    p = doc.add_paragraph(item, style='List Bullet')
    set_no_indent_style(p)

# ========== 3. СИСТЕМНЫЕ ТРЕБОВАНИЯ ==========
system_req_heading = doc.add_heading('3 Системные требования', 1)
set_times_new_roman(system_req_heading)
add_bookmark(system_req_heading, 'bookmark_system_req')

# 3.1 Функциональные требования
func_req_heading = doc.add_heading('3.1 Функциональные требования', 2)
set_times_new_roman(func_req_heading)
add_bookmark(func_req_heading, 'bookmark_func_req')

func_req = [
    ('FR-001', 'Система должна обеспечивать пошаговый игровой процесс с очередностью ходов на основе инициативы юнитов'),
    ('FR-002', 'Система должна поддерживать два игрока, управляющих армиями на одном поле боя'),
    ('FR-003', 'Система должна предоставлять возможность выбора и перемещения юнитов по сетке клеток'),
    ('FR-004', 'Система должна реализовывать механику атаки: ближний бой для ближнебойных юнитов и дальний бой для дальнобойных'),
    ('FR-005', 'Система должна поддерживать различные типы юнитов с уникальными характеристиками (здоровье, атака, защита, скорость, инициатива)'),
    ('FR-006', 'Система должна реализовывать систему отрядов, где каждый отряд имеет количество юнитов и общее здоровье'),
    ('FR-007', 'Система должна поддерживать героев с уникальными характеристиками (атака, защита, знание, сила заклинаний, мана)'),
    ('FR-008', 'Система должна реализовывать систему заклинаний для героев с различными эффектами (урон, исцеление, баффы, дебаффы)'),
    ('FR-009', 'Система должна отслеживать перезарядку заклинаний и расход маны'),
    ('FR-010', 'Система должна реализовывать физическую и магическую атаку и защиту'),
    ('FR-011', 'Система должна поддерживать сопротивление магии для защиты от дебаффов'),
    ('FR-012', 'Система должна определять победителя при уничтожении всех юнитов противника'),
    ('FR-013', 'Система должна отображать визуальную информацию о юнитах (полоски здоровья, тултипы, окна информации)'),
    ('FR-014', 'Система должна воспроизводить звуковые эффекты для действий (атаки, заклинания, смерть)'),
    ('FR-015', 'Система должна поддерживать анимации для действий юнитов (атака, перемещение, заклинания)'),
    ('FR-016', 'Система должна реализовывать режим креатива для размещения юнитов на поле боя'),
    ('FR-017', 'Система должна поддерживать различные фракции (люди, нежить, эльфы, демоны, гномы, тени)'),
    ('FR-018', 'Система должна реализовывать систему удачи и боевого духа, передаваемых от героя к юнитам'),
    ('FR-019', 'Система должна поддерживать контратаки юнитов при получении урона'),
    ('FR-020', 'Система должна реализовывать режим защиты для юнитов, увеличивающий защиту на ход'),
    ('FR-021', 'Система должна поддерживать сохранение и загрузку настроек игры через JSON файлы'),
    ('FR-022', 'Система должна предоставлять режим отладки с визуализацией игровой информации (F1-F6)'),
    ('FR-023', 'Система должна поддерживать веб-версию для запуска в браузере'),
    ('FR-024', 'Система должна поддерживать нативную версию (EXE) для Windows')
]

for req_id, req_desc in func_req:
    p = doc.add_paragraph()
    p.add_run(f'{req_id}: ')
    p.add_run(req_desc)
    set_gost_style(p)

# Дополнительная информация о заклинаниях
spells_heading = doc.add_heading('3.1.1 Описание системы заклинаний', 3)
set_times_new_roman(spells_heading)
add_bookmark(spells_heading, 'bookmark_spells')

spells_info = doc.add_paragraph('Система поддерживает более 30 различных заклинаний, разделенных на категории:')
set_gost_style(spells_info)

spells_categories = [
    ('Боевые заклинания', 'Наносят урон вражеским юнитам: Огненный шар, Молния, Огненная стрела, Ледяная стрела, Кольцо холода, Каменные шипы, Цепная молния, Шок земли, Метеоритный дождь'),
    ('Заклинания исцеления', 'Восстанавливают здоровье союзным юнитам: Лечение, Воскрешение, Поднятие мёртвых'),
    ('Защитные заклинания', 'Увеличивают защиту или создают барьеры: Каменная кожа, Огненный щит, Ледяной щит, Руна защиты, Руна стены'),
    ('Баффы', 'Улучшают характеристики союзных юнитов: Благословение, Ускорение, Руна скорости, Руна магии, Руна берсерка, Точность, Молитва'),
    ('Дебаффы', 'Ухудшают характеристики вражеских юнитов: Проклятие, Замедление, Слабость, Забвение, Ослепление'),
    ('Контроль', 'Ограничивают действия противника: Зыбучие пески, Огненная стена'),
    ('Специальные', 'Уникальные эффекты: Снятие чар, Контрудар, Фантом')
]

for cat_name, cat_desc in spells_categories:
    p = doc.add_paragraph()
    p.add_run(f'{cat_name}: ')
    p.add_run(cat_desc)
    set_gost_style(p)

spells_note = doc.add_paragraph()
spells_note.add_run('Примечание: ')
spells_note.add_run('Заклинания имеют стоимость маны, дальность применения и могут иметь перезарядку. '
                    'Дебаффы могут быть отражены сопротивлением магии цели.')
set_gost_style(spells_note)

# Дополнительная информация о юнитах
units_heading = doc.add_heading('3.1.2 Описание системы юнитов', 3)
set_times_new_roman(units_heading)
add_bookmark(units_heading, 'bookmark_units')

units_info = doc.add_paragraph('Система поддерживает более 40 различных типов юнитов, разделенных по фракциям:')
set_gost_style(units_info)

factions = [
    ('Люди', 'Крестьянин, Копейщик, Арбалетчик, Мечник, Грифон, Кавалерист, Монах, Ангел'),
    ('Нежить', 'Скелет, Зомби, Призрак, Вампир, Лич, Рыцарь Смерти, Костяной Дракон, Жнец'),
    ('Эльфы', 'Пикси, Эльфийский Разведчик, Эльфийский Лучник, Дриада, Энт, Друид, Единорог'),
    ('Демоны', 'Бес, Гог, Демон, Цербер, Суккуб, Дьявол, Адский Конь, Мантикора, Красный Дракон'),
    ('Гномы', 'Шахтер, Метатель Копий, Медвежий Всадник, Рунический Маг, Ярл, Разведчик, Зверь, Минотавр, Кузнечный Дракон, Владыка Гор'),
    ('Тени', 'Ведьма, Ящер-Всадник, Зеленый Дракон, Кровавая Жрица, Смотрящий, Волхв')
]

for faction_name, faction_units in factions:
    p = doc.add_paragraph()
    p.add_run(f'{faction_name}: ')
    p.add_run(faction_units)
    set_gost_style(p)

units_chars = doc.add_paragraph('Каждый юнит имеет следующие характеристики:')
set_gost_style(units_chars)

units_chars_list = [
    'Здоровье (health) — текущее и максимальное здоровье',
    'Физическая атака (phys_attack) — урон физическими атаками',
    'Магическая атака (magic_attack) — урон магическими атаками',
    'Физическая защита (phys_defense) — защита от физических атак',
    'Магическая защита (magic_defense) — защита от магических атак',
    'Сопротивление магии (magic_resist) — процент отражения дебаффов',
    'Скорость (speed) — количество клеток перемещения за ход',
    'Инициатива (initiative) — определяет порядок ходов',
    'Дальнобойность (is_ranged) — возможность атаковать на расстоянии',
    'Количество в отряде (squad_count) — количество юнитов в отряде'
]

for item in units_chars_list:
    p = doc.add_paragraph(item, style='List Bullet')
    set_no_indent_style(p)

# 3.2 Нефункциональные требования
nonfunc_heading = doc.add_heading('3.2 Нефункциональные требования', 2)
set_times_new_roman(nonfunc_heading)
add_bookmark(nonfunc_heading, 'bookmark_nonfunc')

# 3.2.1 Атрибуты качества
quality_heading = doc.add_heading('3.2.1 Атрибуты качества', 3)
set_times_new_roman(quality_heading)
add_bookmark(quality_heading, 'bookmark_quality')

quality_attrs = [
    ('Надёжность', 
     'Система должна корректно обрабатывать все игровые ситуации без критических ошибок. '
     'Игра должна сохранять стабильность при длительной работе. Важность: высокая, так как '
     'ошибки могут привести к потере игрового прогресса. Измерение: количество ошибок на час игры, '
     'процент успешных завершений игровых сессий.'),
    
    ('Производительность',
     'Система должна обеспечивать плавную работу с частотой кадров не менее 30 FPS на целевых устройствах. '
     'Задержка при выполнении действий не должна превышать 100 мс. Важность: средняя, так как игра '
     'не требует высокой производительности, но должна быть отзывчивой. Измерение: FPS во время игры, '
     'время отклика на действия пользователя.'),
    
    ('Удобство использования',
     'Интерфейс должен быть интуитивно понятным, не требующим изучения инструкций для базового использования. '
     'Важность: высокая, так как определяет доступность игры для широкой аудитории. Измерение: время, '
     'необходимое новому пользователю для выполнения базовых действий, количество обращений к справке.'),
    
    ('Совместимость',
     'Система должна работать на Windows 10/11 (нативная версия) и в современных браузерах (веб-версия). '
     'Важность: высокая, так как определяет доступность игры. Измерение: процент успешных запусков '
     'на различных конфигурациях.'),
    
    ('Масштабируемость',
     'Система должна поддерживать добавление новых юнитов, заклинаний и механик без изменения базовой архитектуры. '
     'Важность: средняя, так как позволяет расширять игру в будущем. Измерение: время, необходимое для '
     'добавления нового юнита или заклинания.'),
    
    ('Поддерживаемость',
     'Код должен быть структурирован, документирован и легко читаем для будущих модификаций. '
     'Важность: средняя для конечных пользователей, высокая для разработчиков. Измерение: время, '
     'необходимое разработчику для понимания и модификации кода.')
]

for attr_name, attr_desc in quality_attrs:
    p = doc.add_paragraph()
    p.add_run(f'{attr_name}: ')
    p.add_run(attr_desc)
    set_gost_style(p)
    doc.add_paragraph()

# ========== 4. ЗАКЛЮЧЕНИЕ ==========
conclusion_heading = doc.add_heading('4 Заключение', 1)
set_times_new_roman(conclusion_heading)
add_bookmark(conclusion_heading, 'bookmark_conclusion')

conclusion = doc.add_paragraph(
    'В данном документе были изложены требования к проекту игры "Фэнтези Стратегия" — пошаговой '
    '2D стратегической игры в фэнтези сеттинге. Документ содержит описание функциональных и '
    'нефункциональных требований, характеристик пользователей, программных интерфейсов и системных требований.')
set_gost_style(conclusion)

conclusion2 = doc.add_paragraph(
    'Игра реализует комплексную систему тактических боевых действий с поддержкой более 40 типов юнитов, '
    'более 30 различных заклинаний, системы отрядов, физической и магической боевой системы. Продукт '
    'предназначен для локального использования одним или двумя игроками на одном устройстве.')
set_gost_style(conclusion2)

conclusion3 = doc.add_paragraph(
    'Система разработана с использованием современных технологий (Python, Pygame) и поддерживает '
    'несколько платформ: нативную версию для Windows и веб-версию для браузеров. Все требования, '
    'изложенные в документе, направлены на обеспечение качественного игрового опыта и удобства использования.')
set_gost_style(conclusion3)

conclusion4 = doc.add_paragraph(
    'Документ может служить основой для разработки, тестирования и сопровождения системы, а также '
    'для оценки соответствия реализованного продукта заявленным требованиям.')
set_gost_style(conclusion4)

# Убеждаемся, что весь текст использует Times New Roman 14 и черный цвет
# Убираем жирный текст везде, кроме заголовков
print("Проверка шрифтов и соответствия ГОСТу...")
for paragraph in doc.paragraphs:
    is_heading = paragraph.style.name.startswith('Heading')
    
    # Для заголовков - оставляем жирными, но проверяем шрифт
    if is_heading:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
            run.bold = True  # Заголовки должны быть жирными по ГОСТу
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    else:
        # Для всех остальных параграфов убираем жирный текст
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
            run.bold = False  # Обычный текст не жирный
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# Сохраняем документ
output_path = r'D:\лабы\ЖЦРПО\1 сем\Game_Requirements_Document.docx'
try:
    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    print(f'\n✓ Документ успешно создан по ГОСТу: {output_path}')
    print(f'  Размер файла: {file_size / 1024:.1f} KB')
    print(f'\n' + '='*60)
    print('ПРОВЕРКА СООТВЕТСТВИЯ ГОСТу:')
    print('='*60)
    print('✓ Поля документа: верхнее/нижнее 2 см, левое 3 см, правое 1.5 см')
    print('✓ Шрифт: Times New Roman, размер 14 пт')
    print('✓ Межстрочный интервал: 1.5')
    print('✓ Отступ первой строки: 1.25 см')
    print('✓ Цвет текста: черный')
    print('✓ Заголовки: Times New Roman 14, жирный (Bold)')
    print('✓ Обычный текст: Times New Roman 14, обычный (не жирный)')
    print('✓ Содержание: список с гиперссылками и номерами страниц, Times New Roman 14, черный')
    print('✓ Нумерация страниц: добавлена в нижний колонтитул, Times New Roman 14')
    print('✓ Таблицы: Times New Roman 14, обычный текст')
    print('✓ Списки: Times New Roman 14, обычный текст')
    print('='*60)
    print(f'\nПримечание: Номера страниц в содержании автоматически обновятся при открытии документа в Word')
    print(f'  Если номера не отображаются, нажмите Ctrl+A и F9 для обновления всех полей')
    sys.stdout.flush()
except Exception as e:
    print(f'\n✗ Ошибка при сохранении: {e}')
    import traceback
    traceback.print_exc()
    sys.exit(1)

